import sqlite3
import datetime
import requests
import telebot

from docx import Document

import keyboards
import config
import text


bot = telebot.TeleBot(config.TELEGRAM_TOKEN)


def is_in_database(user_id):
    """Checks if user already in database."""

    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    users = cursor.execute(f'''SELECT COUNT(id) 
                            FROM users 
                            WHERE user_id=?
                            ''', (user_id,)).fetchall()[0][0]
    
    cursor.close()
    database.close()

    return users


def add_user(user_id, user_username):
    """Adds a new user to database."""

    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    cursor.execute(f'''
        INSERT INTO users (user_id, user_username, history)
        VALUES (?, ?, ?)
        ''', (user_id, user_username, "[]",))
        
    database.commit()
    cursor.close()
    database.close()


def add_answer(question, answer, quality):
    """Adds a new user to database."""

    database = sqlite3.connect("db.db", detect_types=sqlite3.PARSE_DECLTYPES|sqlite3.PARSE_COLNAMES)
    cursor = database.cursor()

    timestamp = datetime.datetime.utcnow() + datetime.timedelta(hours=3)

    cursor.execute(f'''
        INSERT INTO {quality}_answers (question, answer, timestamp)
        VALUES (?, ?, ?)
        ''', (question, answer, timestamp,))
        
    database.commit()
    cursor.close()
    database.close()


def select_all_answers(quality):
    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    answers = cursor.execute(f'''SELECT question, answer, timestamp
                            FROM {quality}_answers 
                            ''').fetchall()
    
    cursor.close()
    database.close()

    return answers


def select_new_answers(quality):
    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    answers = cursor.execute(f'''SELECT question, answer, timestamp
                            FROM {quality}_answers
                            WHERE in_report=? 
                            ''', (False,)).fetchall()
    
    cursor.close()
    database.close()

    return answers


def set_in_report(quality):
    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    cursor.execute(f'''UPDATE {quality}_answers
                    SET in_report=?
                    ''', (True,))

    database.commit()
    cursor.close()
    database.close()


def get_message_id(user_id):
    """Gets user's information."""

    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    message_id = cursor.execute(f'''SELECT message_id
                            FROM users 
                            WHERE user_id=?
                            ''', (user_id,)).fetchall()
    
    cursor.close()
    database.close()

    if message_id:
        message_id = message_id[0][0]

    return message_id


def update_message_id(user_id, message_id):
    '''Updates message's id.'''

    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    cursor.execute(f'''UPDATE users
                    SET message_id=?
                    WHERE user_id=?
                    ''', (message_id, user_id,))

    database.commit()
    cursor.close()
    database.close()


def get_info(user_id):
    """Gets user's information."""

    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    info = cursor.execute(f'''SELECT history, busy
                            FROM users 
                            WHERE user_id=?
                            ''', (user_id,)).fetchall()[0]
    
    cursor.close()
    database.close()

    return info


def update_history(user_id, new_history):
    """Updates counter in database."""

    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    cursor.execute(f'''UPDATE users
                    SET history=?
                    WHERE user_id=?
                    ''', (new_history, user_id,))

    database.commit()
    cursor.close()
    database.close()


def set_busy(user_id, status):
    """Updates busy status to True."""

    database = sqlite3.connect("db.db")
    cursor = database.cursor()

    cursor.execute(f'''UPDATE users
                    SET busy=?
                    WHERE user_id=?
                    ''', (status, user_id,))

    database.commit()
    cursor.close()
    database.close()


def connect_ai(change_id, question, message_to_reply, history, user_id):
    """Connects to AI, handles query."""

    headers = {
        'Connection': 'keep-alive',
        'Content-Type': 'application/json',
        'User-Agent': 'python-requests/2.31.0',
        'accept': 'application/json',
        'token': config.AI_TOKEN,
    }

    json_data = {
        'input_text' : question,
        'chat_history' : history,
    }

    try:
        response = requests.post(config.AI_ENDPOINT, headers=headers, json=json_data, timeout=30)

        if response.status_code == 200:
            answer = response.json().get('answer')
            sources = response.json().get('sources')
            
            if sources:
                all_sources = set()
                for source in sources:
                    all_sources.add(source.get('url').split('.')[-2])
            else:
                sources = ''

            all_sources = ', '.join(list(all_sources))

            history_question = {
                'sent' : True,
                'message' : question,
            }

            history_answer= {
                'sent' : False,
                'message' : answer,
            }

            history.append(history_question)
            history.append(history_answer)

            if text.CANT_ANSWER in answer.lower():
                add_answer(question, answer, 'bad')

        else:
            answer = False
            history = False

    except:
        answer = False
        history = False


    if answer and history:
        update_history(user_id, str(history))

        for i in range(1, 20):
            answer = answer.replace(f'{i}.', f'\n\n{i}.')

        answer += f'\n\nИсточники: {all_sources}'
        
        try:
            bot.delete_message(chat_id=user_id,
                               message_id=change_id,
                               )
            
            bot.send_message(chat_id=user_id,
                            text=answer,
                            reply_markup=keyboards.vote_keyboard(),
                            reply_to_message_id=message_to_reply
                            )
        except:
            pass
    
    else:
        try:
            bot.delete_message(chat_id=user_id,
                               message_id=change_id,
                               )
            
            bot.send_message(chat_id=user_id,
                            text=text.ERROR_TEXT,
                            parse_mode='Markdown',
                            )
        except:
            pass
        
    
    set_busy(user_id, False)


def handle_new_report_request(user_id, quality):
    answers_info = select_new_answers(quality)
    set_in_report(quality)

    if answers_info:
        doc = Document()

        for answer in answers_info:
            doc.add_paragraph(f'{answer[2]}\nВОПРОС:\n{answer[0]}\nОТВЕТ:\n{answer[1]}\n\n')
        
        doc.save(f'{quality}.docx')

        with open(f'{quality}.docx', "rb") as f:
            file_data = f.read()

        bot.send_document(chat_id=user_id,
                        document=file_data,
                        visible_file_name=f'{quality}.docx', 
                        )
    else:
        bot.send_message(chat_id=user_id,
                         text=text.NO_DATA,
                         )

    config.REPORT_FLAG = False


def handle_all_report_request(user_id, quality):
    answers_info = select_all_answers(quality)
    set_in_report(quality)

    if answers_info:
        doc = Document()

        for answer in answers_info:
            doc.add_paragraph(f'{answer[2]}\nВОПРОС:\n{answer[0]}\nОТВЕТ:\n{answer[1]}\n\n')
        
        doc.save(f'{quality}.docx')

        with open(f'{quality}.docx', "rb") as f:
            file_data = f.read()

        bot.send_document(chat_id=user_id,
                        document=file_data,
                        visible_file_name=f'{quality}.docx', 
                        )
    else:
        bot.send_message(chat_id=user_id,
                         text=text.NO_DATA,
                         )

    config.REPORT_FLAG = False
